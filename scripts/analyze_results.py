import json
import matplotlib.pyplot as plt
from collections import defaultdict
from docx import Document
from docx.shared import Inches
import os

def analyze_results(results_file, output_folder="analysis_output"):
    # Load the results JSON file
    with open(results_file, 'r') as f:
        data = json.load(f)

    # Separate results into CPU and GPU run types
    cpu_results = defaultdict(lambda: defaultdict(list))
    gpu_results = defaultdict(lambda: defaultdict(list))

    # Organize data into CPU and GPU structures
    for entry in data:
        test_name_parts = entry["Test Name"].split("_")
        test_method = test_name_parts[0]  # Method name (e.g., numba, pythran, clean)
        test_name = "_".join(test_name_parts[1:])  # Test name after the method (e.g., convolution_test)
        run_type = entry["Run Type"]

        # Store in the appropriate section based on run type
        if run_type == "CPU":
            cpu_results[test_name][test_method].append({
                "machine_specs": {
                    "OS": entry["OS"],
                    "CPU": entry["CPU"],
                    "GPU": entry["GPU"].split(",")[0]  # Only take the GPU model
                },
                "time": entry["Elapsed Time"],
                "cache_usage": entry.get("Average Cache Usage (MB)", 0),
                "ram_usage": entry.get("Average Memory Usage (%)", 0),
                "cpu_usage": entry.get("Average CPU Usage (%)", 0),
                "learn_params": {
                    "Workers": entry.get("Workers", None),
                    "Batch Size": entry.get("Batch Size", None),
                    "Learning rate": entry.get("Learning rate", None),
                    "Epochs": entry.get("Epochs", None)
                }
            })
        elif run_type == "GPU":
            gpu_results[test_name][test_method].append({
                "machine_specs": {
                    "OS": entry["OS"],
                    "CPU": entry["CPU"],
                    "GPU": entry["GPU"].split(",")[0]  # Only take the GPU model
                },
                "time": entry["Elapsed Time"],
                "cache_usage": entry.get("Average Cache Usage (MB)", 0),
                "ram_usage": entry.get("Average Memory Usage (%)", 0),
                "cpu_usage": entry.get("Average CPU Usage (%)", 0),
                "learn_params": {
                    "Workers": entry.get("Workers", None),
                    "Batch Size": entry.get("Batch Size", None),
                    "Learning rate": entry.get("Learning rate", None),
                    "Epochs": entry.get("Epochs", None)
                }
            })

    # Create the output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Create a Word document for the report
    doc = Document()
    doc.add_heading("Analysis Report: CPU and GPU Tests", 0)

    # Add page breaks for better organization
    def add_page_break():
        doc.add_page_break()

    # Helper function to plot and save graphs
    def plot_and_save_graph(x_values, y_values, x_label, y_label, title, filename):
        plt.figure()
        plt.plot(x_values, y_values, marker='o', color='blue')
        plt.xlabel(x_label)
        plt.ylabel(y_label)
        plt.title(title)
        plt.tight_layout()
        plt.savefig(filename)
        plt.close()

    # Function to add graphs side by side in the document
    def insert_images_side_by_side(images, doc, width=Inches(2.5)):
        table = doc.add_table(rows=1, cols=len(images))
        table.autofit = True
        for i, image in enumerate(images):
            cell = table.cell(0, i)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image, width=width)

    # Function to generate graphs for both CPU and GPU sections
    def generate_section(test_results, section_name):
        doc.add_heading(section_name, level=1)

        for test_name, methods in test_results.items():
            doc.add_heading(f"Test: {test_name}", level=2)

            for method, entries in methods.items():
                workers_values = [entry["learn_params"]["Workers"] for entry in entries if entry["learn_params"]["Workers"] is not None]
                batch_size_values = [entry["learn_params"]["Batch Size"] for entry in entries if entry["learn_params"]["Batch Size"] is not None]
                learning_rate_values = [entry["learn_params"]["Learning rate"] for entry in entries if entry["learn_params"]["Learning rate"] is not None]
                epochs_values = [entry["learn_params"]["Epochs"] for entry in entries if entry["learn_params"]["Epochs"] is not None]

                elapsed_times = [entry["time"] for entry in entries]
                cache_usages = [entry["cache_usage"] for entry in entries]
                ram_usages = [entry["ram_usage"] for entry in entries]
                cpu_usages = [entry["cpu_usage"] for entry in entries]

                # Generate graphs based on MNIST parameters (only for MNIST)
                if "mnist" in test_name.lower():
                    # Workers vs Performance Graphs
                    if workers_values:
                        # Generate images and save to temp files
                        images = []
                        graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_{method}_workers_vs_time.png")
                        plot_and_save_graph(workers_values, elapsed_times, "Workers", "Elapsed Time (seconds)", f"Elapsed Time vs Workers for {test_name} ({method})", graph_filename)
                        images.append(graph_filename)

                        graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_{method}_workers_vs_cache.png")
                        plot_and_save_graph(workers_values, cache_usages, "Workers", "Cache Usage (MB)", f"Cache Usage vs Workers for {test_name} ({method})", graph_filename)
                        images.append(graph_filename)

                        # Insert images side by side in the document
                        insert_images_side_by_side(images, doc)

                        # Add page break for next graph section
                        add_page_break()

                    # Batch Size vs Performance Graphs
                    if batch_size_values:
                        images = []
                        graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_{method}_batch_size_vs_time.png")
                        plot_and_save_graph(batch_size_values, elapsed_times, "Batch Size", "Elapsed Time (seconds)", f"Elapsed Time vs Batch Size for {test_name} ({method})", graph_filename)
                        images.append(graph_filename)

                        graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_{method}_batch_size_vs_ram.png")
                        plot_and_save_graph(batch_size_values, ram_usages, "Batch Size", "RAM Usage (%)", f"RAM Usage vs Batch Size for {test_name} ({method})", graph_filename)
                        images.append(graph_filename)

                        # Insert images side by side
                        insert_images_side_by_side(images, doc)

                        # Add page break for next graph section
                        add_page_break()

                    # Learning Rate vs Performance Graphs
                    if learning_rate_values:
                        images = []
                        graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_{method}_learning_rate_vs_time.png")
                        plot_and_save_graph(learning_rate_values, elapsed_times, "Learning Rate", "Elapsed Time (seconds)", f"Elapsed Time vs Learning Rate for {test_name} ({method})", graph_filename)
                        images.append(graph_filename)

                        # Insert images side by side
                        insert_images_side_by_side([graph_filename], doc)

                        # Add page break for next graph section
                        add_page_break()

                    # Epochs vs Performance Graphs
                    if epochs_values:
                        images = []
                        graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_{method}_epochs_vs_time.png")
                        plot_and_save_graph(epochs_values, elapsed_times, "Epochs", "Elapsed Time (seconds)", f"Elapsed Time vs Epochs for {test_name} ({method})", graph_filename)
                        images.append(graph_filename)

                        # Insert images side by side
                        insert_images_side_by_side([graph_filename], doc)

                        # Add page break for next section
                        add_page_break()

                # For non-MNIST tests, graph based on methods only
                else:
                    method_names = []
                    avg_times = []
                    avg_cache_usages = []
                    avg_ram_usages = []
                    avg_cpu_usages = []

                    # Collect metrics across methods
                    for method, entries in methods.items():
                        method_names.append(method)
                        avg_time = sum(entry["time"] for entry in entries) / len(entries)
                        avg_cache_usage = sum(entry["cache_usage"] for entry in entries) / len(entries)
                        avg_ram_usage = sum(entry["ram_usage"] for entry in entries) / len(entries)
                        avg_cpu_usage = sum(entry["cpu_usage"] for entry in entries) / len(entries)

                        avg_times.append(avg_time)
                        avg_cache_usages.append(avg_cache_usage)
                        avg_ram_usages.append(avg_ram_usage)
                        avg_cpu_usages.append(avg_cpu_usage)

                    # Elapsed Time
                    images = []
                    graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_elapsed_time.png")
                    plot_and_save_graph(method_names, avg_times, "Methods", "Elapsed Time (seconds)",
                                        f"Elapsed Time for {test_name}", graph_filename)
                    images.append(graph_filename)

                    # Cache Usage
                    graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_cache_usage.png")
                    plot_and_save_graph(method_names, avg_cache_usages, "Methods", "Cache Usage (MB)",
                                        f"Cache Usage for {test_name}", graph_filename)
                    images.append(graph_filename)

                    # RAM Usage
                    graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_ram_usage.png")
                    plot_and_save_graph(method_names, avg_ram_usages, "Methods", "RAM Usage (%)",
                                        f"RAM Usage for {test_name}", graph_filename)
                    images.append(graph_filename)

                    # CPU Usage
                    graph_filename = os.path.join(output_folder, f"{section_name}_{test_name}_cpu_usage.png")
                    plot_and_save_graph(method_names, avg_cpu_usages, "Methods", "CPU Usage (%)",
                                        f"CPU Usage for {test_name}", graph_filename)
                    images.append(graph_filename)

                    # Insert the images side by side and add a page break for the next test
                    insert_images_side_by_side(images, doc)
                    add_page_break()

    # Generate CPU section
    generate_section(cpu_results, "CPU")

    # Generate GPU section
    generate_section(gpu_results, "GPU")

    # Save the document
    output_docx = os.path.join(output_folder, "analysis_report.docx")
    doc.save(output_docx)

    print(f"Analysis complete! Report saved to {output_docx}")

# Clean the previous output and run analysis
os.system("rm -rf EngFinalProject/analysis_output")
analyze_results("EngFinalProject/results/results.json", "EngFinalProject/analysis_output")
